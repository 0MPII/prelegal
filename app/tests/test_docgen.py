import json

import pytest
from docx import Document

from app import docgen

CATALOG = json.loads((docgen.ROOT / "catalog.json").read_text())
CATALOG_FILENAMES = [entry["filename"] for entry in CATALOG]


# --- slugify ---


@pytest.mark.parametrize(
    "text,expected",
    [
        ("Effective Date", "effective_date"),
        ("Governing Law", "governing_law"),
        ("  Leading/Trailing  ", "leading_trailing"),
        ("Provider's", "provider_s"),
        ("Multi   Space", "multi_space"),
        ("ALL CAPS", "all_caps"),
    ],
)
def test_slugify(text, expected):
    assert docgen.slugify(text) == expected


def test_slugify_has_no_leading_or_trailing_underscore():
    slug = docgen.slugify("  ??Weird--Text!!  ")
    assert not slug.startswith("_")
    assert not slug.endswith("_")


# --- role_of ---


@pytest.mark.parametrize("text", ["Provider", "Customer", "Partner", "Company"])
def test_role_of_matches_bare_role(text):
    assert docgen.role_of(text) == text


@pytest.mark.parametrize(
    "text,expected_role",
    [
        ("Provider's", "Provider"),
        ("Provider’s", "Provider"),  # curly apostrophe, as used in some templates
        ("Customer's", "Customer"),
        ("Customer’s", "Customer"),
        ("Partner's", "Partner"),
        ("Company's", "Company"),
    ],
)
def test_role_of_matches_possessive_form(text, expected_role):
    assert docgen.role_of(text) == expected_role


@pytest.mark.parametrize(
    "text",
    ["Effective Date", "Provider Covered Claim", "Purpose", "", "providers", "Providers"],
)
def test_role_of_does_not_match_non_role_text(text):
    assert docgen.role_of(text) is None


# --- build_schema ---


@pytest.mark.parametrize("filename", CATALOG_FILENAMES)
def test_build_schema_succeeds_for_every_catalog_document(filename):
    schema = docgen.build_schema(filename, "Title")
    assert schema.filename == filename
    assert len(schema.roles) >= 2, "every generated document needs at least two parties"
    assert len(set(schema.roles)) == len(schema.roles), "roles should not repeat"


@pytest.mark.parametrize("filename", CATALOG_FILENAMES)
def test_all_field_keys_has_no_duplicates(filename):
    schema = docgen.build_schema(filename, "Title")
    keys = schema.all_field_keys()
    assert len(keys) == len(set(keys))


def test_nda_falls_back_to_generic_two_party_roles():
    # Mutual-NDA.md never references parties by role name inline, unlike
    # the Provider/Customer-style documents, so it must fall back to a
    # generic pair of roles rather than ending up with zero parties.
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    assert schema.roles == ["Party 1", "Party 2"]
    labels = {f.label for f in schema.value_fields}
    assert {"Purpose", "Effective Date", "Governing Law"} <= labels


def test_csa_discovers_roles_and_groups_fields_by_placeholder_class():
    schema = docgen.build_schema("CSA.md", "CSA")
    assert schema.roles == ["Customer", "Provider"]

    groups = dict(schema.value_field_groups())
    assert "Key terms" in groups
    assert "Order form" in groups
    assert any(f.label == "Effective Date" for f in groups["Key terms"])
    assert any(f.label == "Subscription Period" for f in groups["Order form"])


def test_party_field_groups_marks_only_name_as_required():
    schema = docgen.build_schema("CSA.md", "CSA")
    role, fields = schema.party_field_groups()[0]
    assert role == "Customer"

    by_key = {key: (label, placeholder, required) for key, label, placeholder, required in fields}
    assert by_key["customer_name"][2] is True
    assert by_key["customer_address"][2] is False
    assert by_key["customer_signatory_name"][2] is False


# --- fill_spans ---


def test_fill_spans_substitutes_value_field_with_highlight():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    md = docgen.load_source_markdown("Mutual-NDA.md")
    filled = docgen.fill_spans(md, schema, {"purpose": "Evaluating a partnership"})
    assert '<span class="filled-field">Evaluating a partnership</span>' in filled


def test_fill_spans_marks_missing_value_field_as_empty_placeholder():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    md = docgen.load_source_markdown("Mutual-NDA.md")
    filled = docgen.fill_spans(md, schema, {})
    assert '<span class="empty-field">[Purpose]</span>' in filled


def test_fill_spans_unwraps_role_span_but_never_substitutes_the_party_name():
    schema = docgen.build_schema("CSA.md", "CSA")
    md = docgen.load_source_markdown("CSA.md")
    assert '<span class="coverpage_link">Customer</span>' in md  # sanity: source really wraps it

    filled = docgen.fill_spans(md, schema, {"customer_name": "Acme Inc."})
    assert '<span class="coverpage_link">Customer</span>' not in filled
    assert "Acme Inc." not in filled  # the defined term stays put; only the cover table gets the name
    assert "Customer" in filled


def test_fill_spans_ignores_unknown_field_keys():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    md = docgen.load_source_markdown("Mutual-NDA.md")
    # should not raise even if the caller passes keys the schema doesn't know about
    docgen.fill_spans(md, schema, {"not_a_real_field": "x"})


# --- strip_spans_to_text ---


def test_strip_spans_to_text_removes_all_span_tags():
    html = 'Hello <span class="filled-field">World</span> and <span class="empty-field">[X]</span>.'
    assert docgen.strip_spans_to_text(html) == "Hello World and [X]."


# --- build_docx ---


@pytest.mark.parametrize("filename", CATALOG_FILENAMES)
def test_build_docx_produces_a_valid_document_for_every_catalog_entry(filename):
    schema = docgen.build_schema(filename, "Doc Title")
    values = {key: f"value-for-{key}" for key in schema.all_field_keys()}
    buffer = docgen.build_docx(schema, values)

    doc = Document(buffer)
    assert doc.paragraphs[0].text == "Doc Title"
    assert len(doc.tables) == 1


def test_build_docx_cover_table_contains_party_and_value_field_rows():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    values = {
        "party_1_name": "Acme Inc.",
        "party_2_name": "Beta LLC",
        "purpose": "Evaluating a partnership",
        "governing_law": "Delaware",
    }
    buffer = docgen.build_docx(schema, values)
    doc = Document(buffer)

    rows = {row.cells[0].text: row.cells[1].text for row in doc.tables[0].rows}
    assert rows["Party 1"] == "Acme Inc."
    assert rows["Party 2"] == "Beta LLC"
    assert rows["Purpose"] == "Evaluating a partnership"
    assert rows["Governing Law"] == "Delaware"


def test_build_docx_signatory_row_omits_stray_comma_when_blank():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    buffer = docgen.build_docx(schema, {})
    doc = Document(buffer)

    rows = {row.cells[0].text: row.cells[1].text for row in doc.tables[0].rows}
    assert rows["Party 1 signatory"] == ""


def test_build_docx_body_has_no_duplicate_standard_terms_heading():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    buffer = docgen.build_docx(schema, {})
    doc = Document(buffer)

    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert heading_texts.count("Standard Terms") == 1


def test_build_docx_strips_markdown_link_syntax_but_keeps_link_text():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    buffer = docgen.build_docx(schema, {})
    doc = Document(buffer)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[Version 1.0]" not in full_text
    assert "(https://" not in full_text
    assert "Version 1.0" in full_text


def test_build_docx_marks_unfilled_body_blanks_as_bracketed_placeholders():
    schema = docgen.build_schema("Mutual-NDA.md", "Mutual NDA")
    buffer = docgen.build_docx(schema, {})
    doc = Document(buffer)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[Purpose]" in full_text


def test_build_docx_role_terms_stay_in_body_instead_of_the_actual_party_name():
    schema = docgen.build_schema("CSA.md", "CSA")
    values = {"customer_name": "Acme Inc.", "provider_name": "CloudCo"}
    buffer = docgen.build_docx(schema, values)
    doc = Document(buffer)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Customer" in full_text
    assert "Acme Inc." not in full_text  # party name only belongs in the cover table
