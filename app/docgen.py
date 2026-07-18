import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = ROOT / "templates"

LINK_RE = re.compile(r'<span class="([a-z_0-9]+_link)"[^>]*>(.*?)</span>')
BOLD_RE = re.compile(r"\*\*(.*?)\*\*")
MD_LINK_RE = re.compile(r"\[([^\]]*)\]\([^)]*\)")
ROLE_RE = re.compile(r"^(Provider|Customer|Partner|Company)(?:['’]s)?$")

GROUP_LABELS = {
    "coverpage_link": "Deal terms",
    "keyterms_link": "Key terms",
    "orderform_link": "Order form",
    "businessterms_link": "Business terms",
    "sow_link": "Statement of work",
}

PARTY_SUBFIELDS = [
    ("name", "name", "e.g. Acme Inc.", True),
    ("address", "address", "e.g. 123 Market St, San Francisco, CA", False),
    ("signatory_name", "signatory name", "e.g. Jordan Lee", False),
    ("signatory_title", "signatory title", "e.g. CEO", False),
    ("email", "email", "e.g. jordan@acme.com", False),
]


def slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.strip().lower()).strip("_")


def role_of(text: str) -> str | None:
    match = ROLE_RE.match(text.strip())
    return match.group(1) if match else None


@dataclass(frozen=True)
class ValueField:
    key: str
    label: str
    span_text: str
    group: str


@dataclass(frozen=True)
class DocumentSchema:
    filename: str
    title: str
    roles: list[str]
    value_fields: list[ValueField]

    def party_field_groups(self):
        groups = []
        for role in self.roles:
            role_key = slugify(role)
            fields = [
                (f"{role_key}_{suffix}", f"{role} {label}", placeholder, required)
                for suffix, label, placeholder, required in PARTY_SUBFIELDS
            ]
            groups.append((role, fields))
        return groups

    def value_field_groups(self):
        groups: dict[str, list[ValueField]] = {}
        order: list[str] = []
        for f in self.value_fields:
            if f.group not in groups:
                groups[f.group] = []
                order.append(f.group)
            groups[f.group].append(f)
        return [(g, groups[g]) for g in order]

    def all_field_keys(self) -> list[str]:
        keys = []
        for _, fields in self.party_field_groups():
            keys.extend(key for key, *_ in fields)
        keys.extend(f.key for f in self.value_fields)
        return keys


def load_source_markdown(filename: str) -> str:
    return (TEMPLATES_DIR / filename).read_text()


def build_schema(filename: str, title: str) -> DocumentSchema:
    md_text = load_source_markdown(filename)
    roles: list[str] = []
    seen_roles: set[str] = set()
    value_fields: list[ValueField] = []
    seen_values: set[str] = set()

    for match in LINK_RE.finditer(md_text):
        link_class, text = match.group(1), match.group(2)
        role = role_of(text)
        if role:
            if role not in seen_roles:
                seen_roles.add(role)
                roles.append(role)
            continue
        if text not in seen_values:
            seen_values.add(text)
            value_fields.append(
                ValueField(
                    key=slugify(text),
                    label=text,
                    span_text=text,
                    group=GROUP_LABELS.get(link_class, link_class),
                )
            )

    if not roles:
        roles = ["Party 1", "Party 2"]

    return DocumentSchema(filename=filename, title=title, roles=roles, value_fields=value_fields)


def fill_spans(md_text: str, schema: DocumentSchema, values: dict[str, str]) -> str:
    value_by_span = {f.span_text: values.get(f.key, "").strip() for f in schema.value_fields}

    def repl(match: re.Match) -> str:
        text = match.group(2)
        if role_of(text):
            return text
        val = value_by_span.get(text, "")
        if val:
            return f'<span class="filled-field">{val}</span>'
        return f'<span class="empty-field">[{text}]</span>'

    return LINK_RE.sub(repl, md_text)


def strip_spans_to_text(md_text: str) -> str:
    return re.sub(r"</?span[^>]*>", "", md_text)


def _add_runs(paragraph, text: str) -> None:
    parts = BOLD_RE.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1


def build_docx(schema: DocumentSchema, values: dict[str, str]) -> BytesIO:
    md_text = fill_spans(load_source_markdown(schema.filename), schema, values)
    plain_text = strip_spans_to_text(md_text)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(schema.title, level=0)
    doc.add_heading("Cover Page", level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    def add_row(label: str, val: str) -> None:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = val

    for role, fields in schema.party_field_groups():
        role_key = slugify(role)
        signatory_name = values.get(f"{role_key}_signatory_name", "").strip()
        signatory_title = values.get(f"{role_key}_signatory_title", "").strip()
        signatory = ", ".join(p for p in (signatory_name, signatory_title) if p)

        add_row(role, values.get(f"{role_key}_name", ""))
        add_row(f"{role} address", values.get(f"{role_key}_address", ""))
        add_row(f"{role} signatory", signatory)
        add_row(f"{role} email", values.get(f"{role_key}_email", ""))

    for field in schema.value_fields:
        add_row(field.label, values.get(field.key, ""))

    doc.add_paragraph()

    for line in plain_text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        line = MD_LINK_RE.sub(r"\1", line)
        paragraph = doc.add_paragraph()
        _add_runs(paragraph, line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
